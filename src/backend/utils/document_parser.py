import io
import re
import csv
import json
import logging
from typing import List

logger = logging.getLogger(__name__)

def parse_pdf(content: bytes) -> str:
    """Extracts text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
        
        extracted_text = "\n\n".join(text_parts).strip()
        if not extracted_text:
            logger.warning("PDF extraction returned empty text.")
            return "[Empty PDF file]"
        return extracted_text
    except Exception as e:
        logger.error(f"Error parsing PDF file: {e}", exc_info=True)
        raise ValueError(f"Failed to parse PDF document: {str(e)}")

def parse_docx(content: bytes) -> str:
    """Extracts text from a DOCX file using python-docx."""
    try:
        from docx import Document
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        text_parts = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    # Remove duplicate cell contents if cells are merged
                    cleaned_row = []
                    for val in row_text:
                        if not cleaned_row or cleaned_row[-1] != val:
                            cleaned_row.append(val)
                    text_parts.append(" | ".join(cleaned_row))
                    
        extracted_text = "\n".join(text_parts).strip()
        if not extracted_text:
            logger.warning("DOCX extraction returned empty text.")
            return "[Empty DOCX file]"
        return extracted_text
    except Exception as e:
        logger.error(f"Error parsing DOCX file: {e}", exc_info=True)
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")

def parse_doc_fallback(content: bytes) -> str:
    """Fallback text extractor for legacy .doc binary documents or other binary text types."""
    try:
        # Extract sequences of printable characters of length 4 or more
        pattern = re.compile(rb'[a-zA-Z0-9\s.,!?;:"\'()\[\]\-\/\x80-\xff]{4,}')
        chunks = pattern.findall(content)
        
        lines = []
        for chunk in chunks:
            try:
                text = chunk.decode('utf-8', errors='ignore').strip()
                # Ensure it has readable content and at least some alphabetic letters
                if len(text) > 4 and any(c.isalpha() for c in text):
                    # Clean up multiple whitespaces
                    text = re.sub(r'\s+', ' ', text)
                    lines.append(text)
            except Exception:
                pass
                
        extracted_text = "\n".join(lines).strip()
        if not extracted_text:
            return content.decode('utf-8', errors='replace')
        return extracted_text
    except Exception as e:
        logger.error(f"Error in fallback DOC parser: {e}")
        return content.decode('utf-8', errors='replace')

def parse_csv(content: bytes) -> str:
    """Parses and formats CSV content into text representation."""
    try:
        decoded = content.decode("utf-8")
    except UnicodeDecodeError:
        decoded = content.decode("utf-8", errors="replace")
        
    try:
        reader = csv.reader(io.StringIO(decoded))
        rows = list(reader)
        formatted_rows = []
        for row in rows:
            cleaned_row = [col.strip() for col in row]
            formatted_rows.append(", ".join(cleaned_row))
        return "\n".join(formatted_rows).strip()
    except Exception as e:
        logger.error(f"Error formatting CSV: {e}")
        return decoded

def parse_json(content: bytes) -> str:
    """Parses and formats JSON content into formatted indented string."""
    try:
        decoded = content.decode("utf-8")
    except UnicodeDecodeError:
        decoded = content.decode("utf-8", errors="replace")
        
    try:
        obj = json.loads(decoded)
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error(f"Error parsing JSON: {e}")
        return decoded

def parse_document(content: bytes, filename: str) -> str:
    """Main entry point to parse a document's binary content based on its file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    
    if ext == "pdf":
        return parse_pdf(content)
    elif ext == "docx":
        return parse_docx(content)
    elif ext == "doc":
        return parse_doc_fallback(content)
    elif ext == "csv":
        return parse_csv(content)
    elif ext == "json":
        return parse_json(content)
    elif ext in ("txt", "md"):
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("utf-8", errors="replace")
    else:
        # Catch-all fallback
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("utf-8", errors="replace")
